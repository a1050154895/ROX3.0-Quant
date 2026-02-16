from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any, Tuple, Set
import os
import json
import re
import traceback
import logging

HAS_NLP = False
cosine_similarity = None
logger = logging.getLogger(__name__)

@dataclass
class KnowledgeDocument:
    path: str
    title: str
    content: str
    vector: Optional[List[float]] = field(default=None)
    category: Optional[str] = field(default=None)
    author: Optional[str] = field(default=None)
    keywords: List[str] = field(default_factory=list)
    concepts: List[str] = field(default_factory=list)

class KnowledgeBase:
    def __init__(self):
        self.documents: List[KnowledgeDocument] = []
        self.model = None
        self._nlp_checked = False
        self._category_mapping = {
            "economics": ["资本论", "就业、利息和货币通论", "经济学", "宏观", "微观"],
            "investment": ["穷查理宝典", "原则", "投资学", "笑傲股市", "教你炒股票"],
            "policy": ["结构性改革", "置身事内", "规划纲要", "政策"],
            "history": ["南明史", "乡土中国", "历史", "文化"],
            "philosophy": ["传习录", "乌合之众", "自卑与超越", "哲学"]
        }

    def _ensure_nlp(self) -> bool:
        global HAS_NLP
        global cosine_similarity
        if self._nlp_checked:
            return bool(HAS_NLP and self.model and cosine_similarity)
        self._nlp_checked = True
        try:
            from sentence_transformers import SentenceTransformer
            from sklearn.metrics.pairwise import cosine_similarity as _cos
            cosine_similarity = _cos
            HAS_NLP = True
            self.model = SentenceTransformer("all-MiniLM-L6-v2")
            logger.info("✓ SentenceTransformer loaded successfully")
            return True
        except ImportError:
            HAS_NLP = False
            self.model = None
            logger.warning("⚠️  SentenceTransformer not installed")
            return False
        except Exception as e:
            logger.error(f"Failed to load SentenceTransformer: {e}")
            HAS_NLP = False
            self.model = None
            return False

    def _normalize_text(self, s: str) -> str:
        lines = [re.sub(r"\s+", " ", ln).strip() for ln in s.splitlines()]
        lines = [ln for ln in lines if ln]
        return "\n".join(lines)

    def _parse_pdf(self, fp: str) -> str:
        content = ""
        try:
            from pdfminer.high_level import extract_text as _pdf_extract  # type: ignore
            content = _pdf_extract(fp) or ""
        except Exception:
            content = ""
        if not content:
            try:
                import fitz  # type: ignore
                d = fitz.open(fp)
                content = "\n".join([p.get_text() or "" for p in d])
            except Exception:
                content = ""
        if not content:
            try:
                from PyPDF2 import PdfReader  # type: ignore
                r = PdfReader(fp)
                content = "\n".join([(page.extract_text() or "") for page in r.pages])
            except Exception:
                content = ""
        return self._normalize_text(content or "")
    
    def _compute_embedding(self, text: str) -> Optional[List[float]]:
        if not self._ensure_nlp():
            return None
        try:
            # 截断过长的文本以适应模型限制
            embedding = self.model.encode(text[:1000])
            return embedding.tolist()
        except Exception as e:
            logger.error(f"Embedding error: {e}")
            return None

    def _detect_category(self, title: str) -> Optional[str]:
        """自动检测书籍类别"""
        for category, keywords in self._category_mapping.items():
            for keyword in keywords:
                if keyword in title:
                    return category
        return None

    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        try:
            import jieba
            from jieba.analyse import extract_tags
            return extract_tags(text, topK=10)
        except Exception:
            return []

    def build_embedded_from_dir(self, src_dir: str) -> int:
        arr = []
        if not os.path.isdir(src_dir):
            return 0
        
        logger.info(f"Building KB from {src_dir}...")
        processed_count = 0
        
        for root, _, files in os.walk(src_dir):
            for f in files:
                if f.startswith("~$") or f.startswith("."):
                    continue
                fp = os.path.join(root, f)
                ext = os.path.splitext(fp)[1].lower()
                title = os.path.splitext(os.path.basename(fp))[0]
                content = ""
                try:
                    if ext in [".txt", ".md"]:
                        with open(fp, "r", encoding="utf-8", errors="ignore") as fh:
                            content = self._normalize_text(fh.read())
                    elif ext == ".docx":
                        try:
                            from docx import Document  # type: ignore
                            d = Document(fp)
                            content = self._normalize_text("\n".join([p.text for p in d.paragraphs]))
                        except Exception:
                            content = ""
                    elif ext == ".pdf":
                        content = self._parse_pdf(fp)
                    else:
                        content = ""
                except Exception as e:
                    logger.error(f"Error processing {title}: {e}")
                    content = ""
                
                if content:
                    # 生成向量
                    vec = self._compute_embedding(title + "\n" + content)
                    category = self._detect_category(title)
                    keywords = self._extract_keywords(content)
                    
                    arr.append({
                        "path": fp, 
                        "title": title, 
                        "content": content, 
                        "vector": vec,
                        "category": category,
                        "keywords": keywords
                    })
                    logger.info(f"Processed: {title} (Category: {category})")
                    processed_count += 1

        base = os.path.dirname(__file__)
        assets_dir = os.path.join(base, "assets")
        if not os.path.exists(assets_dir):
            os.makedirs(assets_dir)
            
        out = os.path.join(assets_dir, "embedded_kb.json")
        try:
            with open(out, "w", encoding="utf-8") as f:
                json.dump(arr, f, ensure_ascii=False)
            logger.info(f"Saved {len(arr)} docs to {out}")
        except Exception as e:
            logger.error(f"Failed to save KB: {e}")
            return 0
        return processed_count

    def load_embedded(self) -> int:
        self.documents.clear()
        # Use resource_utils if available, otherwise fallback
        try:
            from ..resource_utils import get_resource_path
            path = get_resource_path(os.path.join("app", "rox_quant", "assets", "embedded_kb.json"))
        except ImportError:
             base = os.path.dirname(__file__)
             path = os.path.join(base, "assets", "embedded_kb.json")

        if not os.path.isfile(path):
            # Try direct relative path if resource path fails or returns non-existent
             base = os.path.dirname(__file__)
             local_path = os.path.join(base, "assets", "embedded_kb.json")
             if os.path.isfile(local_path):
                 path = local_path
             else:
                 logger.warning("No embedded KB found")
                 return 0

        try:
            with open(path, "r", encoding="utf-8") as f:
                arr = json.load(f)
            if isinstance(arr, list):
                for item in arr:
                    t = str(item.get("title") or "").strip()
                    c = str(item.get("content") or "").strip()
                    p = str(item.get("path") or "")
                    v = item.get("vector") # List[float] or None
                    cat = item.get("category")
                    keywords = item.get("keywords", [])
                    
                    if t and c:
                        doc = KnowledgeDocument(
                            path=p, 
                            title=t, 
                            content=self._normalize_text(c), 
                            vector=v,
                            category=cat,
                            keywords=keywords
                        )
                        self.documents.append(doc)
            logger.info(f"✓ Loaded {len(self.documents)} documents")
        except Exception as e:
            logger.error(f"Failed to load KB: {e}")
            self.documents = []
        return len(self.documents)

    def load_dir(self, path: str) -> int:
        self.documents.clear()
        if not os.path.isdir(path):
            return 0
        
        processed_count = 0
        for root, _, files in os.walk(path):
            for f in files:
                if f.startswith("~$") or f.startswith("."):
                    continue
                fp = os.path.join(root, f)
                ext = os.path.splitext(fp)[1].lower()
                title = os.path.splitext(os.path.basename(fp))[0]
                content = ""
                try:
                    if ext in [".txt", ".md"]:
                        with open(fp, "r", encoding="utf-8", errors="ignore") as fh:
                            content = self._normalize_text(fh.read())
                    elif ext == ".docx":
                        try:
                            from docx import Document  # type: ignore
                            d = Document(fp)
                            content = self._normalize_text("\n".join([p.text for p in d.paragraphs]))
                        except Exception:
                            content = ""
                    elif ext == ".pdf":
                        content = self._parse_pdf(fp)
                    else:
                        content = ""
                except Exception as e:
                    logger.error(f"Error processing {title}: {e}")
                    content = ""
                
                if content:
                    vec = self._compute_embedding(title + "\n" + content)
                    category = self._detect_category(title)
                    keywords = self._extract_keywords(content)
                    
                    doc = KnowledgeDocument(
                        path=fp, 
                        title=title, 
                        content=content, 
                        vector=vec,
                        category=category,
                        keywords=keywords
                    )
                    self.documents.append(doc)
                    processed_count += 1
        
        logger.info(f"Loaded {processed_count} documents from {path}")
        return processed_count

    def size(self) -> int:
        return len(self.documents)

    def search(self, query: str, limit: int = 5, category: Optional[str] = None) -> List[KnowledgeDocument]:
        # 混合搜索：如果有向量则结合语义，否则仅关键词
        q = query.strip().lower()
        if not q:
            return []
        
        # 过滤文档
        filtered_docs = self.documents
        if category:
            filtered_docs = [doc for doc in self.documents if doc.category == category]
        
        # 1. 语义搜索 (Semantic Search)
        semantic_scores = {} # id -> score
        if self._ensure_nlp():
            try:
                q_vec = self.model.encode(q)
                # 收集所有有向量的文档
                valid_docs = [(i, d.vector) for i, d in enumerate(filtered_docs) if d.vector]
                if valid_docs:
                    ids = [x[0] for x in valid_docs]
                    vecs = [x[1] for x in valid_docs]
                    
                    # 计算相似度
                    sims = cosine_similarity([q_vec], vecs)[0]
                    
                    for idx, score in zip(ids, sims):
                        semantic_scores[idx] = float(score)
            except Exception as e:
                logger.error(f"Semantic search error: {e}")

        # 2. 关键词搜索 (Keyword Search)
        keyword_scores = {}
        for i, doc in enumerate(filtered_docs):
            s = 0
            if q in doc.title.lower():
                s += 3.0
            if q in doc.content.lower():
                s += 1.0
            # 简单的词频统计
            s += doc.content.lower().count(q) * 0.1
            # 关键词匹配加分
            for keyword in doc.keywords:
                if q in keyword.lower():
                    s += 0.5
            if s > 0:
                keyword_scores[i] = s

        # 3. 融合分数 (Fusion)
        # 归一化关键词分数
        if keyword_scores:
            max_kw = max(keyword_scores.values())
            if max_kw > 0:
                for k in keyword_scores:
                    keyword_scores[k] /= max_kw # map to 0-1

        final_scores = []
        for i, doc in enumerate(filtered_docs):
            sem_s = semantic_scores.get(i, 0.0)
            kw_s = keyword_scores.get(i, 0.0)
            
            # 加权融合: 语义 0.7 + 关键词 0.3 (可调整)
            if HAS_NLP and self.model:
                total = sem_s * 0.7 + kw_s * 0.3
            else:
                total = kw_s
            
            if total > 0.01: # 阈值
                final_scores.append((total, doc))

        final_scores.sort(key=lambda x: x[0], reverse=True)
        return [d for _, d in final_scores[:limit]]

    def get_context_for_algo(self, topic: str) -> str:
        """
        为算法模块提供特定主题的上下文知识
        """
        docs = self.search(topic, limit=3)
        if not docs:
            return ""
        
        context = []
        for d in docs:
            # 截取摘要
            snippet = d.content[:500].replace("\n", " ")
            context.append(f"《{d.title}》: {snippet}...")
        return "\n".join(context)

    def count_sector_keywords(self) -> dict:
        keywords = {
            "新能源": 0,
            "半导体": 0,
            "人工智能": 0,
            "数据资产": 0,
            "光伏": 0,
            "电动车": 0,
            "医药": 0,
            "消费": 0,
            "白酒": 0,
            "低空经济": 0,
            "华为": 0,
            "算力": 0
        }
        for doc in self.documents:
            text = doc.content.lower()
            for k in list(keywords.keys()):
                keywords[k] += text.count(k.lower())
        return keywords

    def count_macro_keywords(self) -> dict:
        keywords = {
            "加息": 0,
            "降息": 0,
            "通胀": 0,
            "GDP": 0,
            "宽货币": 0,
            "紧信用": 0,
            "政策支持": 0,
            "风险": 0,
            "复苏": 0,
            "流动性": 0
        }
        for doc in self.documents:
            text = doc.content.lower()
            for k in list(keywords.keys()):
                keywords[k] += text.count(k.lower())
        return keywords

    def get_documents_by_category(self, category: str) -> List[KnowledgeDocument]:
        """按类别获取文档"""
        return [doc for doc in self.documents if doc.category == category]

    def get_categories(self) -> Set[str]:
        """获取所有类别"""
        categories = set()
        for doc in self.documents:
            if doc.category:
                categories.add(doc.category)
        return categories

    def get_top_keywords(self, limit: int = 20) -> Dict[str, int]:
        """获取所有文档的top关键词"""
        keyword_count = {}
        for doc in self.documents:
            for keyword in doc.keywords:
                keyword_count[keyword] = keyword_count.get(keyword, 0) + 1
        
        # 排序并返回前N个
        sorted_keywords = sorted(keyword_count.items(), key=lambda x: x[1], reverse=True)
        return dict(sorted_keywords[:limit])

    def build_knowledge_graph(self) -> Dict[str, Any]:
        """构建简单的知识图谱"""
        graph = {
            "nodes": [],
            "edges": []
        }
        
        # 添加文档节点
        for i, doc in enumerate(self.documents):
            graph["nodes"].append({
                "id": f"doc_{i}",
                "label": doc.title,
                "type": "document",
                "category": doc.category
            })
        
        # 添加关键词节点并建立连接
        keyword_id_map = {}
        keyword_counter = 0
        
        for i, doc in enumerate(self.documents):
            for keyword in doc.keywords:
                if keyword not in keyword_id_map:
                    keyword_id_map[keyword] = f"kw_{keyword_counter}"
                    graph["nodes"].append({
                        "id": keyword_id_map[keyword],
                        "label": keyword,
                        "type": "keyword"
                    })
                    keyword_counter += 1
                
                graph["edges"].append({
                    "source": f"doc_{i}",
                    "target": keyword_id_map[keyword],
                    "type": "contains"
                })
        
        return graph

    def get_related_documents(self, doc_index: int, limit: int = 5) -> List[Tuple[float, KnowledgeDocument]]:
        """获取相关文档"""
        if doc_index >= len(self.documents):
            return []
        
        target_doc = self.documents[doc_index]
        if not target_doc.vector:
            return []
        
        related = []
        for i, doc in enumerate(self.documents):
            if i == doc_index or not doc.vector:
                continue
            
            try:
                similarity = cosine_similarity([target_doc.vector], [doc.vector])[0][0]
                if similarity > 0.5:
                    related.append((similarity, doc))
            except Exception:
                pass
        
        related.sort(key=lambda x: x[0], reverse=True)
        return related[:limit]